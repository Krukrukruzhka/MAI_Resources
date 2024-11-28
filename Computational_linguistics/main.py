import aiohttp
import asyncio
import glob
import logging
import re

from bs4 import BeautifulSoup
from docx import Document
from word_classes import WORD_CLASSES


logger = logging.getLogger()
logger.setLevel("DEBUG")


def _prepare_word(word: str) -> str:
    letters = []
    for symbol in word:
        if symbol == "ё":
            symbol = "е"
        if "а" <= symbol <= "я" or symbol == "-":
            letters.append(symbol)
    return "".join(letters)


def _prepare_noun_endings(endings: list[str], root: str) -> None:
    logger.warning(f"all forms: {endings}")
    ends = ["ей", "ь", "й", "а", "о", "я", "е", ""]
    i = 0
    for end in ends:
        if endings[i] == root + end:
            endings[i] = end
            break
    else:
        for end in ends:
            if len(endings[i]) >= len(end) and endings[i][-len(end):] == end:
                endings[i] = end
                break
        else:
            endings[i] = ""

    ends = ["ом", "ем", "ей", "ой", "ым", "ю", ""]
    i = 1
    for end in ends:
        if endings[i] == root + end:
            endings[i] = end
            break
    else:
        for end in ends:
            if len(endings[i]) >= len(end) and endings[i][-len(end):] == end:
                endings[i] = end
                break
        else:
            endings[i] = ""

    ends = ["ы", "и", "а", "я", "е", ""]
    i = 2
    for end in ends:
        if endings[i] == root + end:
            endings[i] = end
            break
    else:
        for end in ends:
            if len(endings[i]) >= len(end) and endings[i][-len(end):] == end:
                endings[i] = end
                break
        else:
            endings[i] = ""

    ends = ["ов", "ей", "ев", "ов", "ых", "ий", "ь", "й", ""]
    i = 3
    for end in ends:
        if endings[i] == root + end:
            endings[i] = end
            break
    else:
        for end in ends:
            if len(endings[i]) >= len(end) and endings[i][-len(end):] == end:
                endings[i] = end
                break
        else:
            endings[i] = ""


async def get_word_html(word: str) -> str:
    async with aiohttp.ClientSession() as session:
        url = f"http://ru.wiktionary.org/wiki/{word}"
        logger.warning(f"GET HTML from {url} ")
        async with session.get(url, ssl=False) as resp:
            if resp.status != 200:
                logger.exception(f"Something went wrong with word {word}. Status code {resp.status}.")
                if resp.status == 404:
                    return "слово не нашлось или его не существует"
                return f"HTTP error. Status code {resp.status}"
            return await resp.text()


async def get_word_metadata(word: str, root: str) -> list[str | None]:
    word = word.lower().replace("ё", "е")

    word_html = await get_word_html(word)
    if len(word_html) > 4 and (word_html.startswith("HTTP") or word_html.startswith("слово не нашлось")):
        return [word_html]

    soup = BeautifulSoup(word_html, 'html.parser')
    try:
        div = soup.find("div", {"class": "mw-content-ltr mw-parser-output"})
        p_elements = div.find_all("p")

        word_params = p_elements[1].text.lower().replace("; ", ", ").split(", ")
        for i in range(len(word_params)):
            word_params[i] = word_params[i].replace("ё", "е")

        word_type = word_params[0].replace("\n", "")
        logger.warning(word_type)
        word_params = set(word_params)

        metadata = [word_type]

        if word_type == "существительное":
            if "мужской род" in word_params or "мужской" in word_params:
                metadata.append("мужской")
            elif "женский род" in word_params or "женский" in word_params or "общий род (может согласовываться с другими частями речи как мужского" in word_params:
                metadata.append("женский")
            elif "средний род" in word_params or "средний" in word_params:
                metadata.append("средний")
            else:
                raise RuntimeError("Слово без рода")

            if "одушевленное" in word_params:
                metadata.append("одушевленный")
            elif "неодушевленное" in word_params:
                metadata.append("неодушевленный")
            else:
                raise RuntimeError(r"Без души ¯\_(ツ)_/¯")

            raws = div.find("table", {"class": "morfotable ru"}).find_all("tr")

            end1 = raws[1].find_all("td")[1].get_text(separator=" ").strip().split()[0]
            end2 = raws[5].find_all("td")[1].get_text(separator=" ").strip().split()[0]

            end3 = raws[1].find_all("td")[2]
            if end3.get_text(separator=" ").strip().split()[0] == "*":
                end3 = end3.find("span").get_text(separator=" ").strip().split()[0]
            else:
                end3 = end3.get_text(separator=" ").strip().split()[0]

            end4 = raws[2].find_all("td")[2]
            if end4.get_text(separator=" ").strip().split()[0] == "*":
                end4 = end4.find("span").get_text(separator=" ").strip().split()[0]
            else:
                end4 = end4.get_text(separator=" ").strip().split()[0]

            endings = [
                _prepare_word(end1),
                _prepare_word(end2),
                _prepare_word(end3),
                _prepare_word(end4)
            ]

            _prepare_noun_endings(endings, root)
            metadata.append(" ".join(endings))

        elif word_type == "глагол":
            metadata = ["глагол"]
        elif word_type == "числительное":
            if len(word) > 2 and word[-3:] == "еро":
                metadata = ["еро"]
            elif word[-1] == "ь":
                metadata = ["ь"]
            else:
                metadata.append(word)
        elif word_type in ("деепричастие", "причастие", "наречие", "предикатив"):
            metadata = [word_type]
        else:
            metadata = ["Unknown type"]
        return metadata

    except Exception as exept:
        return [f"Ошибка {exept}"]


async def identity_word_class(root: str, end: str) -> str:
    tmp = WORD_CLASSES
    word = root + end
    metadata = await get_word_metadata(word, root)
    logger.warning(metadata)

    if metadata and len(metadata[0]) > 6 and (metadata[0].startswith("Ошибка") or metadata[0].startswith("HTTP") or metadata[0].startswith("слово не нашлось")):
        logger.warning(metadata[0])
        logger.warning("\n\n")
        return metadata[0]

    for param in metadata:
        tmp = tmp.get(param, None)
        if tmp is None:
            logger.warning(f"word: {word}\t\tparams: {metadata}")
            logger.warning("Ошибка Invalid WORD_CLASSES structure\n\n")
            return "Ошибка Invalid WORD_CLASSES structure"

    logger.warning(tmp)
    logger.warning("\n\n")
    return tmp


async def generate_answer_file():
    path = f"res/input/*.docx"
    docx_files = glob.glob(path)

    for file_path in docx_files:
        doc = Document(file_path)
        filename = file_path.split("/")[-1]

        output_file_path = f"res/output/{filename[:-5]}.txt"
        with open(output_file_path, 'w+', encoding='utf-8') as output_file:
            coruts = []
            for paragraph in doc.paragraphs:
                text = re.sub(r'\.', '', re.sub(r'\s+', ' ', paragraph.text))
                text_elements = text.split()
                if 3 <= len(text_elements):
                    root, end = text_elements[0], (text_elements[1] if text_elements[1] != "+" else "")
                    root, end = _prepare_word(root), _prepare_word(end)
                    coruts.append(asyncio.create_task(identity_word_class(root, end)))

            coruts = await asyncio.gather(*coruts)
            i = 0
            for paragraph in doc.paragraphs:
                text = re.sub(r'\.', '', re.sub(r'\s+', ' ', paragraph.text))
                text_elements = text.split()
                if 3 <= len(text_elements):
                    root, end = text_elements[0], (text_elements[1] if text_elements[1] != "+" else "")
                    root, end = _prepare_word(root), _prepare_word(end)
                    new_class = coruts[i]
                    if new_class == "062" and root[-1] == "и":
                        new_class = "061"
                    if len(text_elements) == 3 and new_class == text_elements[-1].split('/')[-1]:
                        text_elements.append("success")
                    else:
                        text_elements.append(f"#{new_class}")
                    output_file.write(" ".join(text_elements) + '\n')
                    i += 1


if __name__ == "__main__":
    asyncio.run(generate_answer_file())
    # print(asyncio.run(identity_word_class("брам-стеньг", "а")))
